#!/usr/bin/env python3
"""Convert any Markdown file to a formatted Word document (.docx).

Usage:
    python md_to_docx.py architecture.md   → writes architecture.docx
"""

import re
import sys
from pathlib import Path

try:
    from PIL import ImageFont as _ImageFont
    _PIL_OK = True
except ImportError:
    _PIL_OK = False
    _ImageFont = None  # type: ignore[assignment]

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Constants ────────────────────────────────────────────────────────────────
FONT = "Calibri"
BODY_PT = Pt(12)
H2_PT = Pt(14)   # ## and #
H3_PT = Pt(13)   # ###
MARGIN = Cm(2.5)

# A4 portrait: 21 cm − 2×2.5 cm margins = 16 cm usable width
PAGE_USABLE_CM = 21.0 - 2 * 2.5

# A4 usable height ≈ 29.7 − 5 = 24.7 cm; at 12pt single-spacing ≈ 0.5 cm/line
TABLE_WARN_ROWS = 45

# ── Emoji / zero-width strip ──────────────────────────────────────────────────
# Explicit ranges so en dash (U+2013) and em dash (U+2014) are NOT stripped.
_EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"   # misc symbols, pictographs, transport, etc.
    "\U00002600-\U000027BF"   # misc symbols, dingbats
    "︀-️"           # variation selectors
    "​-‍"           # zero-width space / non-joiner / joiner
    "﻿"                  # BOM / zero-width no-break space
    "]+",
    flags=re.UNICODE,
)

# ── PIL font loading for accurate column-width measurement ────────────────────
_FONT_PT = 12           # matches Word BODY_PT (12pt)
_PX_TO_CM = 2.54 / 72  # PIL renders at 72 DPI; 1 px = 1 pt = 1/72 inch

_FONT_REG_PATHS = [
    # Calibri — not installed by default on macOS
    "/Library/Fonts/Calibri.ttf",
    "/Library/Fonts/Microsoft/Calibri.ttf",
    "C:/Windows/Fonts/calibri.ttf",
    # Carlito — metric-compatible open-source substitute for Calibri
    "/usr/share/fonts/truetype/crosextra/Carlito-Regular.ttf",
    # Arial — closest widely available substitute on macOS / Linux
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial.ttf",
    "C:/Windows/Fonts/arial.ttf",
    "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
]
_FONT_BOLD_PATHS = [
    "/Library/Fonts/Calibri Bold.ttf",
    "/Library/Fonts/Microsoft/Calibri Bold.ttf",
    "C:/Windows/Fonts/calibrib.ttf",
    "/usr/share/fonts/truetype/crosextra/Carlito-Bold.ttf",
    "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    "/Library/Fonts/Arial Bold.ttf",
    "C:/Windows/Fonts/arialbd.ttf",
    "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
]

_FONT_REG = None
_FONT_BOLD = None
_BOLD_IS_FALLBACK = False  # True when bold font not found; regular used + 1.08× scale
_FONT_SOURCE = "not loaded"
_FONT_BOLD_SOURCE = "not loaded"
_FONTS_LOADED = False


def _load_fonts() -> None:
    global _FONT_REG, _FONT_BOLD, _BOLD_IS_FALLBACK
    global _FONT_SOURCE, _FONT_BOLD_SOURCE, _FONTS_LOADED
    if _FONTS_LOADED:
        return
    _FONTS_LOADED = True
    if not _PIL_OK:
        _FONT_SOURCE = _FONT_BOLD_SOURCE = "PIL not available"
        return
    for path in _FONT_REG_PATHS:
        try:
            _FONT_REG = _ImageFont.truetype(path, _FONT_PT)
            _FONT_SOURCE = path
            break
        except (IOError, OSError):
            pass
    if _FONT_REG is None:
        _FONT_SOURCE = _FONT_BOLD_SOURCE = "no usable TTF found; using char-count fallback"
        return
    for path in _FONT_BOLD_PATHS:
        try:
            _FONT_BOLD = _ImageFont.truetype(path, _FONT_PT)
            _FONT_BOLD_SOURCE = path
            break
        except (IOError, OSError):
            pass
    if _FONT_BOLD is None:
        _FONT_BOLD = _FONT_REG
        _FONT_BOLD_SOURCE = _FONT_SOURCE
        _BOLD_IS_FALLBACK = True
    print(f"Font (regular): {_FONT_SOURCE}")
    bold_note = " [fallback regular + 1.08x]" if _BOLD_IS_FALLBACK else " [true bold]"
    print(f"Font (bold):    {_FONT_BOLD_SOURCE}{bold_note}")


def _measure_text_cm(text: str, *, is_header: bool = False) -> float:
    """Return rendered width of text in cm using PIL font metrics at 12pt."""
    _load_fonts()
    if _FONT_REG is None:
        return len(text) * 0.21
    font = _FONT_BOLD if is_header else _FONT_REG
    cm = font.getlength(text) * _PX_TO_CM
    if is_header and _BOLD_IS_FALLBACK:
        cm *= 1.08
    return cm


# Mirrors _inline_cell break points used for w:noBreak: spaces, @, period, slash, underscore
_TOKEN_SPLIT_RE = re.compile(r"[\s@./_]+")

# Floor measurement only: split "MedicationRequest" → ["Medication", "Request"] so the floor
# reflects each camelCase segment, not the full compound word.  Rendering is unchanged.
_CAMEL_SPLIT_RE = re.compile(r"(?<=[a-z])(?=[A-Z])")


def _clean(text: str) -> str:
    text = _EMOJI_RE.sub("", text)
    # Normalize all dashes to plain hyphen-minus
    text = text.replace("---", "-")
    text = text.replace("--", "-")
    text = text.replace("—", "-")   # em dash
    text = text.replace("–", "-")   # en dash
    return text.strip()


# ── Paragraph helpers ─────────────────────────────────────────────────────────
def _fmt_para(para, keep_with_next: bool = False) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    if keep_with_next:
        pf.keep_with_next = True
    # Suppress auto-hyphenation and spell/grammar check per paragraph
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn("w:suppressAutoHyphens")) is None:
        el = OxmlElement("w:suppressAutoHyphens")
        pPr.append(el)


def _run(para, text: str, *, bold=False, italic=False, size=None, font=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font or FONT
    r.font.size = size or BODY_PT
    # Disable spell/grammar check and proofing on every run
    rPr = r._r.get_or_add_rPr()
    if rPr.find(qn("w:noProof")) is None:
        rPr.append(OxmlElement("w:noProof"))
    return r


# ── Inline Markdown renderer ──────────────────────────────────────────────────
_INLINE_RE = re.compile(
    r"("
    r"\*\*\*.+?\*\*\*"          # ***bold italic***
    r"|___.+?___"                # ___bold italic___
    r"|``.+?``"                  # ``code``
    r"|\*\*.+?\*\*"              # **bold**
    r"|__.+?__"                  # __bold__
    r"|`.+?`"                    # `code`
    r"|\*[^*\n]+?\*"             # *italic*
    r"|(?<!\w)_[^_\n]+?_(?!\w)" # _italic_ (word-boundary guarded)
    r")"
)


def _inline(para, text: str, base_size=None) -> None:
    text = text.replace("—", "-").replace("–", "-")
    text = _clean(text)
    bs = base_size or BODY_PT
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if re.fullmatch(r"\*\*\*.+?\*\*\*|___.+?___", part):
            _run(para, part[3:-3], bold=True, italic=True, size=bs)
        elif re.fullmatch(r"\*\*.+?\*\*|__.+?__", part):
            _run(para, part[2:-2], bold=True, size=bs)
        elif re.fullmatch(r"``.+?``", part):
            _run(para, part[2:-2], size=bs)
        elif re.fullmatch(r"`.+?`", part):
            _run(para, part[1:-1], size=bs)
        elif re.fullmatch(r"\*[^*\n]+?\*|(?<!\w)_[^_\n]+?_(?!\w)", part):
            _run(para, part[1:-1], italic=True, size=bs)
        else:
            _run(para, part, size=bs)


def _plain(text: str) -> str:
    """Strip inline markdown to plain text (used for width estimation)."""
    text = _EMOJI_RE.sub("", text)
    text = re.sub(r"\*\*\*(.+?)\*\*\*|___(.+?)___",
                  lambda m: (m.group(1) or m.group(2)), text)
    text = re.sub(r"\*\*(.+?)\*\*|__(.+?)__",
                  lambda m: (m.group(1) or m.group(2)), text)
    text = re.sub(r"``(.+?)``|`(.+?)`",
                  lambda m: (m.group(1) or m.group(2)), text)
    text = re.sub(r"\*([^*\n]+?)\*|(?<!\w)_([^_\n]+?)_(?!\w)",
                  lambda m: (m.group(1) or m.group(2)), text)
    return text.strip()


# ── Table helpers ─────────────────────────────────────────────────────────────
def _cant_split_rows(table) -> None:
    """Mark every row cantSplit so table rows never break across pages."""
    for row in table.rows:
        tr = row._tr
        tr_pr = tr.find(qn("w:trPr"))
        if tr_pr is None:
            tr_pr = OxmlElement("w:trPr")
            tr.insert(0, tr_pr)
        if tr_pr.find(qn("w:cantSplit")) is None:
            cs = OxmlElement("w:cantSplit")
            cs.set(qn("w:val"), "true")
            tr_pr.append(cs)


def _no_borders(tbl_el) -> None:
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tbl_pr.append(borders)


def _apply_col_widths(table, widths_cm: list) -> None:
    tbl = table._tbl

    # Build tblGrid
    tbl_grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(Cm(w).twips)))
        tbl_grid.append(gc)

    # Remove existing tblGrid, insert after tblPr
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        total_twips = sum(int(Cm(w).twips) for w in widths_cm)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(total_twips))
        tblW.set(qn("w:type"), "dxa")
        tbl_pr.append(tblW)

        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tbl_pr.append(tblLayout)

        tbl_pr.addnext(tbl_grid)
    else:
        tbl.insert(0, tbl_grid)

    # Per-cell widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_cm):
                break
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                tc.insert(0, tc_pr)
            for old in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(Cm(widths_cm[i]).twips)))
            tcW.set(qn("w:type"), "dxa")
            tc_pr.append(tcW)


def _parse_table_rows(lines: list) -> list:
    """Return list of cell-lists; separator rows are dropped."""
    rows = []
    for line in lines:
        s = line.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        cells = [c.strip() for c in s.split("|")]
        if all(re.fullmatch(r"[-: ]+", c) for c in cells if c):
            continue  # separator row
        rows.append(cells)
    return rows


def _col0_entity_test(para, text: str, is_header: bool) -> None:
    """Fill col-0 of an entity test table.

    Splits on '_': each segment run gets w:noBreak so Word cannot break inside
    a segment.  The '_' separator runs do NOT get w:noBreak, so those positions
    remain the only legal line-break opportunities.
    """
    parts = text.split("_")
    for idx, part in enumerate(parts):
        r = para.add_run(part)
        r.font.name = FONT
        r.font.size = BODY_PT
        if is_header:
            r.bold = True
        rPr = r._r.get_or_add_rPr()
        rPr.append(OxmlElement("w:noBreak"))
        if rPr.find(qn("w:noProof")) is None:
            rPr.append(OxmlElement("w:noProof"))
        if idx < len(parts) - 1:
            sep = para.add_run("_")
            sep.font.name = FONT
            sep.font.size = BODY_PT
            if is_header:
                sep.bold = True
            rPr_sep = sep._r.get_or_add_rPr()
            if rPr_sep.find(qn("w:noProof")) is None:
                rPr_sep.append(OxmlElement("w:noProof"))
            zwsp = para.add_run("​")
            zwsp.font.name = FONT
            zwsp.font.size = BODY_PT
            if is_header:
                zwsp.bold = True
            rPr_zw = zwsp._r.get_or_add_rPr()
            if rPr_zw.find(qn("w:noProof")) is None:
                rPr_zw.append(OxmlElement("w:noProof"))


def _inline_cell(para, text: str, is_header: bool = False) -> None:
    """Render table cell text with w:noBreak per word to prevent mid-word line breaks."""
    text = text.replace("—", "-").replace("–", "-")
    text = _clean(text)

    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if re.fullmatch(r"\*\*\*.+?\*\*\*|___.+?___", part):
            seg_text, seg_bold, seg_italic = part[3:-3], True, True
        elif re.fullmatch(r"\*\*.+?\*\*|__.+?__", part):
            seg_text, seg_bold, seg_italic = part[2:-2], True, False
        elif re.fullmatch(r"``.+?``", part):
            seg_text, seg_bold, seg_italic = part[2:-2], is_header, False
        elif re.fullmatch(r"`.+?`", part):
            seg_text, seg_bold, seg_italic = part[1:-1], is_header, False
        elif re.fullmatch(r"\*[^*\n]+?\*|(?<!\w)_[^_\n]+?_(?!\w)", part):
            seg_text, seg_bold, seg_italic = part[1:-1], is_header, True
        else:
            seg_text, seg_bold, seg_italic = part, is_header, False

        for token in re.split(r"( +)", seg_text):
            if not token:
                continue
            if token.strip() and ("@" in token or "." in token or "/" in token or "_" in token):
                for sub in re.split(r"([@./_])", token):
                    if not sub:
                        continue
                    r = para.add_run(sub)
                    r.bold = seg_bold
                    r.italic = seg_italic
                    r.font.name = FONT
                    r.font.size = BODY_PT
                    rPr = r._r.get_or_add_rPr()
                    if rPr.find(qn("w:noProof")) is None:
                        rPr.append(OxmlElement("w:noProof"))
                    if sub not in ("@", ".", "/", "_"):
                        if rPr.find(qn("w:noBreak")) is None:
                            rPr.append(OxmlElement("w:noBreak"))
                    elif sub == "_":
                        zwsp = para.add_run("​")
                        zwsp.bold = seg_bold
                        zwsp.italic = seg_italic
                        zwsp.font.name = FONT
                        zwsp.font.size = BODY_PT
                        rPr_zw = zwsp._r.get_or_add_rPr()
                        if rPr_zw.find(qn("w:noProof")) is None:
                            rPr_zw.append(OxmlElement("w:noProof"))
            else:
                r = para.add_run(token)
                r.bold = seg_bold
                r.italic = seg_italic
                r.font.name = FONT
                r.font.size = BODY_PT
                rPr = r._r.get_or_add_rPr()
                if rPr.find(qn("w:noProof")) is None:
                    rPr.append(OxmlElement("w:noProof"))
                if token.strip():
                    if rPr.find(qn("w:noBreak")) is None:
                        rPr.append(OxmlElement("w:noBreak"))


# ── Known table overrides (Jane-approved widths; values unchanged from prior version)
_KNOWN_COL_WIDTHS: dict = {
    ("Test File", "Tests", "REQ_IDs Covered", "Layer(s)", "Gate Group", "Status"):
        [4.0, 1.5, 3.0, 2.0, 3.5, 2.0],          # Section 1 Overview
    ("Test Function", "TC", "Layer", "What Is Verified"):
        [4.5, 2.0, 2.0, 7.5],                     # Entity test tables
    ("Test Function", "Layer", "What Is Verified"):
        [4.5, 2.0, 9.5],                           # Section 4 DB Layer
    ("Gate Group", "Test Type", "Rationale"):
        [5.0, 3.0, 8.0],                           # Section 6.2
    ("Group", "Current Status", "Phase 2 Gate"):
        [5.0, 3.0, 8.0],                           # Section 6.3
    ("Test File", "Test Function", "TC"):
        [3.5, 10.5, 2.0],                          # Section 6.5
}


def _add_table(doc, tbl_lines: list) -> None:
    rows = _parse_table_rows(tbl_lines)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    if n_rows > TABLE_WARN_ROWS:
        print(
            f"WARNING: Table with {n_rows} data rows likely exceeds one page "
            "– please review manually."
        )

    header_key = tuple(_plain(c) for c in rows[0]) if rows else ()
    col_cm = list(_KNOWN_COL_WIDTHS[header_key]) if header_key in _KNOWN_COL_WIDTHS else None

    if col_cm is None:
        # Universal algorithm — Points 2–4: non-iterative, PIL-measured, guaranteed correct.
        CELL_PAD_CM = 0.25

        col_max_cell = [1] * n_cols    # total char count per column (surplus weight)
        col_eff_floor = [0.0] * n_cols  # max measured unbreakable-token width per column

        for r_idx, row in enumerate(rows):
            is_hdr = (r_idx == 0)
            for i, cell in enumerate(row[:n_cols]):
                plain = _plain(cell)
                col_max_cell[i] = max(col_max_cell[i], len(plain))
                for tok in _TOKEN_SPLIT_RE.split(plain):
                    for sub_tok in _CAMEL_SPLIT_RE.split(tok):
                        if sub_tok:
                            col_eff_floor[i] = max(
                                col_eff_floor[i],
                                _measure_text_cm(sub_tok, is_header=is_hdr),
                            )

        effective_floor = [col_eff_floor[i] + CELL_PAD_CM for i in range(n_cols)]
        floor_total = sum(effective_floor)

        surplus = PAGE_USABLE_CM - floor_total
        if surplus <= 0:
            print(
                f"WARNING: table floor_total {floor_total:.3f} cm >= "
                f"PAGE_USABLE_CM {PAGE_USABLE_CM:.1f} cm "
                f"– floors cannot all be honored; proportional-floor fallback applied. "
                f"Headers: {list(header_key)}"
            )
            col_cm = [effective_floor[i] * PAGE_USABLE_CM / floor_total
                      for i in range(n_cols)]
        else:
            weight_total = sum(col_max_cell) or 1
            col_cm = [
                effective_floor[i] + surplus * col_max_cell[i] / weight_total
                for i in range(n_cols)
            ]

    is_col0_testname = col_cm in ([4.5, 2.0, 2.0, 7.5], [4.5, 2.0, 9.5], [3.5, 10.5, 2.0])

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(table._tbl)
    _apply_col_widths(table, col_cm)
    _cant_split_rows(table)
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        if tbl_pr.find(qn("w:keepLines")) is None:
            tbl_pr.append(OxmlElement("w:keepLines"))
        if tbl_pr.find(qn("w:tblLook")) is None:
            tl = OxmlElement("w:tblLook")
            tl.set(qn("w:val"), "04A0")
            tbl_pr.append(tl)

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        for c_idx in range(n_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.clear()
            _fmt_para(para)
            if is_col0_testname and (c_idx == 0 or (col_cm == [3.5, 10.5, 2.0] and c_idx == 1)):
                _col0_entity_test(para, text, is_header)
            else:
                _inline_cell(para, text, is_header)


# ── Page numbers ──────────────────────────────────────────────────────────────
def _add_page_numbers(doc) -> None:
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.name = FONT
        run.font.size = BODY_PT

        fc_begin = OxmlElement("w:fldChar")
        fc_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fc_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._r.append(instr)

        fc_end = OxmlElement("w:fldChar")
        fc_end.set(qn("w:fldCharType"), "end")
        run._r.append(fc_end)


# ── Table line detection ───────────────────────────────────────────────────────
def _is_table_row(line: str) -> bool:
    """True if the line looks like a markdown table row."""
    s = line.strip()
    return "|" in s and (s.startswith("|") or s.count("|") >= 2)


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    # Remove leading/trailing pipes
    s = s.strip("|")
    return bool(re.fullmatch(r"[\s\-:|]+", s))


# ── Document-level defaults ───────────────────────────────────────────────────
def _suppress_hyphenation_defaults(doc) -> None:
    """Set w:suppressAutoHyphens in docDefaults so it applies document-wide."""
    styles_el = doc.styles.element          # <w:styles>
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    ppr_default = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default is None:
        ppr_default = OxmlElement("w:pPrDefault")
        doc_defaults.append(ppr_default)

    ppr = ppr_default.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        ppr_default.append(ppr)

    if ppr.find(qn("w:suppressAutoHyphens")) is None:
        ppr.append(OxmlElement("w:suppressAutoHyphens"))


# ── Main converter ─────────────────────────────────────────────────────────────
def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Page layout: A4 portrait, 2.5 cm margins
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = MARGIN
    sec.right_margin = MARGIN
    sec.top_margin = MARGIN
    sec.bottom_margin = MARGIN

    _suppress_hyphenation_defaults(doc)

    # Default Normal style baseline
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_PT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    _add_page_numbers(doc)

    # Strip blockquote markers (> or >> etc.) from every line before processing
    _bq = re.compile(r"^(>\s*)+")
    lines = [_bq.sub("", l) for l in md_path.read_text(encoding="utf-8").splitlines()]
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # Blank line
        if not line:
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            for cl in code_lines:
                p = doc.add_paragraph()
                _fmt_para(p)
                _run(p, cl)
            continue

        # Markdown table: current line is a row and next non-blank is a separator
        if _is_table_row(line):
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx < len(lines) and _is_table_separator(lines[next_idx]):
                tbl_lines = []
                while i < len(lines) and _is_table_row(lines[i]):
                    tbl_lines.append(lines[i])
                    i += 1
                anchor = doc.add_paragraph()
                _fmt_para(anchor, keep_with_next=True)
                anchor.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                anchor.paragraph_format.line_spacing = Pt(1)
                _add_table(doc, tbl_lines)
                continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)               # _clean is called inside _inline
            size = H2_PT if level <= 2 else H3_PT if level == 3 else BODY_PT
            bold = level <= 4
            p = doc.add_paragraph()
            _fmt_para(p, keep_with_next=True)
            _inline(p, text, base_size=size)
            if bold:
                for run in p.runs:          # force every run bold (incl. leading numbers)
                    run.bold = True
            i += 1
            continue

        # Horizontal rule — skip
        if re.fullmatch(r"[-*_]{3,}", line):
            i += 1
            continue

        # Bullet list — collect consecutive bullet lines
        if re.match(r"^[-*+] ", line):
            while i < len(lines) and re.match(r"^[-*+] ", lines[i].strip()):
                text = re.sub(r"^[-*+] ", "", lines[i].strip())
                p = doc.add_paragraph()
                _fmt_para(p)
                _run(p, "• ")   # bullet character, not a Word list style
                _inline(p, text)
                i += 1
            continue

        # Numbered list — collect consecutive numbered lines
        if re.match(r"^\d+\.", line):
            while i < len(lines) and re.match(r"^\d+\.", lines[i].strip()):
                m2 = re.match(r"^(\d+)\.\s*(.*)", lines[i].strip())
                label = m2.group(1) if m2 else "1"
                text = m2.group(2) if m2 else lines[i].strip()
                p = doc.add_paragraph()
                _fmt_para(p)
                _run(p, f"{label}. ")
                _inline(p, text)
                i += 1
            continue

        # Regular paragraph — join consecutive body lines (standard MD behaviour)
        parts = []
        while i < len(lines):
            l = lines[i].strip()
            if (
                not l
                or re.match(r"^#{1,6}\s", l)
                or re.match(r"^[-*+] ", l)
                or re.match(r"^\d+\.", l)
                or _is_table_row(l)
                or l.startswith("```")
                or re.fullmatch(r"[-*_]{3,}", l)
            ):
                break
            parts.append(l)
            i += 1
        if parts:
            p = doc.add_paragraph()
            _fmt_para(p)
            _inline(p, " ".join(parts))

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


# ── Entry point ───────────────────────────────────────────────────────────────
def main() -> None:
    if len(sys.argv) != 2:
        print("Usage: python md_to_docx.py <file.md>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: file not found: {md_path}")
        sys.exit(1)
    if md_path.suffix.lower() not in (".md", ".markdown"):
        print(f"Warning: {md_path.name} does not have a .md extension, proceeding anyway.")
    docx_path = md_path.with_suffix(".docx")
    convert(md_path, docx_path)


if __name__ == "__main__":
    main()
